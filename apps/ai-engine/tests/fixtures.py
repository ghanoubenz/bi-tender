"""Generate small fixture documents in code (no binaries in the repo)."""

from __future__ import annotations

ITT_TEXT_PAGE1 = (
    "INVITATION TO TENDER\n\n"
    "Tender Ref: ITT-2026-0042\n\n"
    "Issued by: Ministry of Water Resources\n\n"
    "Project: Construction of a Water Treatment Plant, Sohar, Oman\n"
)

ITT_TEXT_PAGE2 = (
    "SECTION 2 SUBMISSION\n\n"
    "The submission deadline: 15 October 2026 at 12:00 local time.\n\n"
    "Bids shall remain valid for a period of 120 days.\n\n"
    "A tender security of OMR 50,000 must accompany the bid.\n"
)


def make_pdf() -> bytes:
    """Two-page born-digital PDF built with pypdf-writable primitives."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    for text in (ITT_TEXT_PAGE1, ITT_TEXT_PAGE2):
        page = writer.add_blank_page(width=612, height=792)
        # pypdf cannot draw text; embed a minimal content stream manually.
        from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

        lines = []
        y = 750
        for line in text.splitlines():
            safe = line.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
            lines.append(f"BT /F1 12 Tf 50 {y} Td ({safe}) Tj ET")
            y -= 18
        stream = DecodedStreamObject()
        stream.set_data("\n".join(lines).encode("latin-1"))
        stream_ref = writer._add_object(stream)
        page[NameObject("/Contents")] = stream_ref
        font = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
        font_ref = writer._add_object(font)
        resources = DictionaryObject({NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})})
        page[NameObject("/Resources")] = resources

    import io

    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def make_docx() -> bytes:
    import io

    from docx import Document

    doc = Document()
    doc.add_heading("Invitation to Tender", level=1)
    doc.add_paragraph("Tender Ref: ITT-2026-0042")
    doc.add_paragraph("Issued by: Ministry of Water Resources")
    doc.add_heading("Submission", level=2)
    doc.add_paragraph("The submission deadline: 15 October 2026 at 12:00 local time.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Bid validity"
    table.rows[1].cells[1].text = "120 days"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_xlsx() -> bytes:
    import io

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "BoQ"
    ws.append(["Item", "Description", "Qty", "Unit"])
    ws.append(["1.1", "Excavation works", 1200, "m3"])
    ws.append(["1.2", "Concrete C40", 800, "m3"])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def make_requirements_docx() -> bytes:
    """A DOCX carrying real obligations, including a requirements table."""
    import io

    from docx import Document

    doc = Document()
    doc.add_heading("Section 7 Qualification Requirements", level=1)
    doc.add_paragraph(
        "7.3.2 The Bidder shall have completed 3 similar projects in the last 5 years."
    )
    doc.add_paragraph(
        "The Contractor must hold a valid ISO 9001 certification at the time of submission."
    )
    doc.add_paragraph("Mobilization shall be completed within 21 calendar days of contract award.")
    doc.add_paragraph("The region has a hot arid climate with seasonal sandstorms.")
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Ref"
    table.rows[0].cells[1].text = "Requirement"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "The bidder shall provide a valid trade licence copy with the bid."
    table.rows[2].cells[0].text = "2"
    table.rows[2].cells[1].text = "Bidders should provide a cost breakdown for all unit rates."
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
