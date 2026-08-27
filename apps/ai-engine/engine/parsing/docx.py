from __future__ import annotations

import io

from docx import Document as DocxDocument

from engine.parsing.base import ParsedBlock, ParseOutput


def parse_docx(data: bytes) -> ParseOutput:
    doc = DocxDocument(io.BytesIO(data))
    out = ParseOutput()
    section: str | None = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading"):
            section = text
            out.blocks.append(ParsedBlock("heading", text, section_path=section))
        else:
            out.blocks.append(ParsedBlock("paragraph", text, section_path=section))
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        flat = "\n".join(" | ".join(r) for r in rows)
        out.blocks.append(ParsedBlock("table", flat, section_path=section, table_data=rows))
    return out
